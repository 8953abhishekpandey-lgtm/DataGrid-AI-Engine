"""
╔══════════════════════════════════════════════════════════════════════════════╗
║  file_handlers.py  —  DataGrid Intelligence · File Upload Handlers         ║
╠══════════════════════════════════════════════════════════════════════════════╣
║  Is file mein upload ki gayi files ko process karne ka logic hai.          ║
║                                                                             ║
║  ⚠️  SIRF YE FORMATS ACCEPT HOTE HAIN:                                     ║
║    Tabular  : .parquet ONLY (Excel/CSV/JSON NAHI)                          ║
║    Documents: .pdf, .docx, .txt, .md                                       ║
║                                                                             ║
║  Main function: process_upload() — ek file ko process karo aur result dict ║
║  return karo jo app.py mein use hota hai.                                  ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

import re                   # Regular expressions — filename sanitization ke liye
from pathlib import Path     # File paths handle karne ke liye
from typing import Optional, Tuple  # Type hints ke liye

import pandas as pd          # DataFrame operations ke liye

from config import SUPPORTED_FORMATS  # Accepted formats ki dict config se


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: FILENAME HELPER FUNCTIONS
# Safe aur valid filenames aur table names banane ke liye
# ══════════════════════════════════════════════════════════════════════════════

def sanitize_table_name(filename: str) -> str:
    """
    Filename se ek valid DuckDB table naam banata hai.

    Rules:
      - Sirf alphanumeric characters aur underscore allowed hain
      - Special characters replace hote hain underscore se
      - Consecutive underscores ek mein merge hote hain
      - Digit se shuru nahi ho sakta (SQL convention)
      - Maximum 64 characters
      - Lowercase karta hai

    Example:
      "Q1-2024 Meter Data.parquet" → "q1_2024_meter_data"
      "123data.parquet"            → "t_123data"
    """
    name = Path(filename).stem          # extension hatao, sirf filename rakhao
    # Example: "data.parquet" → "data", "report.pdf" → "report"

    name = re.sub(r"[^a-zA-Z0-9_]", "_", name)
    # Saare non-alphanumeric characters (spaces, hyphens, dots) ko _ se replace karo
    # [^...] = NOT in this set, \s, -, . etc. sab _ ban jayenge

    name = re.sub(r"_+", "_", name).strip("_")
    # Multiple consecutive underscores ek mein merge karo: "a__b" → "a_b"
    # .strip("_") = shuru aur end ke underscores hatao

    if not name or name[0].isdigit():
        name = "t_" + name  # Agar khali hai ya digit se shuru hai toh "t_" prefix lagao
        # DuckDB mein identifiers digit se shuru nahi ho sakte

    return name.lower()[:64]
    # .lower() = sab lowercase karo (consistency ke liye)
    # [:64] = maximum 64 characters rakho


def secure_filename(filename: str) -> str:
    """
    Upload ki gayi file ka secure naam return karta hai.
    Path traversal attacks prevent karta hai (e.g., "../../etc/passwd").

    Process:
      1. Sirf filename rakhao (directory components hatao)
      2. Dangerous characters remove karo
      3. Khali ho gaya toh 'unnamed' use karo
    """
    filename = Path(filename).name
    # .name = sirf filename ka last component
    # "../../etc/passwd" → "passwd" (path traversal prevent)

    filename = re.sub(r"[^\w\s\-.]", "", filename).strip()
    # Sirf ye characters allow karo: word chars (a-z,A-Z,0-9,_), spaces, hyphen, dot
    # [^\w\s\-.] = NOT word/space/hyphen/dot → remove karo
    # .strip() = edges se whitespace hatao

    return filename or "unnamed"  # Agar sab kuch remove ho gaya toh "unnamed" use karo


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: TABULAR FILE LOADER
# Sirf Parquet files load hoti hain
# ══════════════════════════════════════════════════════════════════════════════

def load_parquet_file(file_path: Path) -> pd.DataFrame:
    """
    Parquet file ko pandas DataFrame mein load karta hai.

    Parquet format kyun use karte hain:
      - Columnar storage = column-wise queries bahut fast hoti hain
      - Built-in compression = file size kam hota hai
      - Data types preserve hote hain (datetime, int, float sab sahi hota hai)
      - Large datasets ke liye ideal (millions of rows)
      - Industry standard for data pipelines

    Raises ValueError agar file parquet nahi hai.
    """
    ext = file_path.suffix.lower()  # file extension lowercase mein

    if ext == ".parquet":
        return pd.read_parquet(file_path)
        # pandas ka built-in parquet reader
        # pyarrow ya fastparquet backend use karta hai (automatically detect karta hai)
        # Returns: DataFrame with all columns and correct dtypes

    else:
        # Koi aur format aaya toh clear error do
        accepted = ".parquet"  # sirf yahi accept hota hai
        raise ValueError(
            f"❌ Format '{ext}' accept nahi hota.\n"
            f"✅ Sirf {accepted} files upload karo.\n"
            f"   Excel (.xlsx) aur CSV (.csv) files supported NAHI hain.\n"
            f"   Apne data ko pehle Parquet mein convert karo:\n"
            f"   Python: df.to_parquet('output.parquet', index=False)"
        )


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: DOCUMENT EXTRACTORS
# PDF, DOCX, TXT files se text extract karne ke functions
# (Ye AI ke saath discuss ho sakte hain lekin SQL queries nahi chalti)
# ══════════════════════════════════════════════════════════════════════════════

def extract_pdf_text(file_path: Path) -> str:
    """
    PDF file se text extract karta hai.
    Har page ka text "[Page N]" header ke saath include hota hai.

    Requirements: pip install pypdf
    Raises ValueError agar pypdf install nahi hai.
    """
    try:
        from pypdf import PdfReader  # pypdf library import karo
    except ImportError:
        raise ValueError("pypdf install nahi hai. Chalao: pip install pypdf")

    reader = PdfReader(str(file_path))  # PDF file open karo
    pages  = []                         # extracted pages ki list

    for i, page in enumerate(reader.pages):  # har page ke liye
        text = page.extract_text() or ""     # text extract karo (None bhi aa sakta hai)
        if text.strip():                     # agar page mein kuch text hai
            pages.append(f"[Page {i + 1}]\n{text.strip()}")
            # Page number header ke saath add karo

    return "\n\n".join(pages)  # pages ko double newline se join karo


def extract_docx_text(file_path: Path) -> str:
    """
    Word (.docx) file se text extract karta hai.
    Headings ko ## prefix ke saath mark karta hai.
    Tables bhi extract hote hain.

    Requirements: pip install python-docx
    """
    try:
        from docx import Document  # python-docx library
    except ImportError:
        raise ValueError("python-docx install nahi hai. Chalao: pip install python-docx")

    doc   = Document(str(file_path))  # Word file open karo
    parts = []                        # extracted text parts

    # Paragraphs process karo
    for para in doc.paragraphs:
        if not para.text.strip():    # empty paragraph skip karo
            continue

        style = para.style.name if para.style else ""
        # paragraph ka style naam (e.g., "Heading 1", "Normal")

        if "Heading" in style:
            parts.append(f"\n## {para.text.strip()}")  # heading mein ## prefix
        else:
            parts.append(para.text.strip())             # normal paragraph as-is

    # Tables process karo
    for table in doc.tables:
        rows = []
        for row in table.rows:
            # Cells ko pipe (|) se join karo — markdown table format
            rows.append(" | ".join(c.text.strip() for c in row.cells))
        if rows:
            parts.append("\n" + "\n".join(rows))  # table add karo

    return "\n".join(parts)  # sab parts join karo


def extract_txt_text(file_path: Path) -> str:
    """
    Plain text file padh ke return karta hai.
    Multiple encodings try karta hai (UTF-8 fail hone pe Latin-1, CP1252).
    """
    # Encodings try karne ki list (common se rare order mein)
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_path.read_text(encoding=enc)
            # Successful hoga toh return karke function exit ho jaayega
        except UnicodeDecodeError:
            continue  # Ye encoding nahi chali, agla try karo

    # Koi bhi encoding nahi chali toh binary padh ke force decode karo
    return file_path.read_bytes().decode("utf-8", errors="replace")
    # errors="replace" = invalid characters ko replacement character (?) se replace karo


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: MAIN UPLOAD PROCESSOR
# Ye main function hai jo app.py call karta hai
# ══════════════════════════════════════════════════════════════════════════════

def process_upload(file_path: Path, safe_name: str) -> dict:
    """
    Uploaded file ko process karke result dict return karta hai.

    Args:
        file_path: disk pe file ka path
        safe_name: sanitized original filename

    Returns:
        Tabular file ke liye:
            {"type": "table", "name": table_name, "filename": safe_name, "df": DataFrame}
        Document ke liye:
            {"type": "document", "name": doc_name, "filename": safe_name,
             "text": extracted_text, "format": "PDF"/"DOCX"/"TXT"/..., "size": char_count}

    Raises ValueError:
        - Unsupported format
        - File extraction failed
        - Parquet file corrupt hai
    """
    ext = file_path.suffix.lower()   # file extension extract karo
    fmt = SUPPORTED_FORMATS.get(ext)  # config se handler type lo
    # fmt = "tabular", "document", ya None (agar format supported nahi)

    if not fmt:
        # Format supported nahi → clear error
        supported = ", ".join(sorted(SUPPORTED_FORMATS.keys()))
        raise ValueError(
            f"❌ Format '{ext}' support nahi hota.\n"
            f"✅ Supported formats: {supported}\n"
            f"⚠️  Tabular data ke liye SIRF .parquet accept hota hai."
        )

    name = sanitize_table_name(safe_name)  # valid table/doc naam banao

    # ── Tabular files (sirf Parquet) ───────────────────────────────────────────
    if fmt == "tabular":
        df = load_parquet_file(file_path)
        # Ye internally check karta hai ki extension .parquet hai
        # Koi bhi exception raise ho toh caller mein handle hogi

        return {
            "type":     "table",      # yeh ek database table banega
            "name":     name,         # DuckDB mein yeh naam se register hoga
            "filename": safe_name,    # original filename (display ke liye)
            "df":       df,           # DataFrame (register_table() ko pass hoga)
        }

    # ── Document files (PDF, DOCX, TXT, MD) ───────────────────────────────────
    elif fmt == "document":
        if ext == ".pdf":
            text = extract_pdf_text(file_path)   # PDF se text nikalo
        elif ext == ".docx":
            text = extract_docx_text(file_path)  # Word file se text nikalo
        else:
            # .txt, .md, .log — sab plain text readers se padhte hain
            text = extract_txt_text(file_path)   # Text file padho

        return {
            "type":     "document",            # yeh ek document hai, table nahi
            "name":     name,                  # document ka naam
            "filename": safe_name,             # original filename
            "text":     text,                  # extracted text content
            "format":   ext[1:].upper(),       # "PDF", "DOCX", "TXT", "MD"
            # ext[1:] = extension ke pehle dot ko hatao, .upper() = uppercase
            "size":     len(text),             # character count (UI mein dikhata hai)
        }

    # ── Unknown format (should never reach here given earlier check) ───────────
    raise ValueError(f"Unknown format handler: {fmt}")
    # Safety net — agar config mein naya format add kiya lekin handler nahi likha